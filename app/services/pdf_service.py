"""
PDF report generator — Pro tier only.

Bilingual (ខ្មែរ / English) using Khmer OS fonts (Cambodian standard).

Fonts used:
  KhmerMuol       (KhmerOSMuolLight.ttf)  — headers / decorative titles
  KhmerBattambang (KhmerOSBattambang.ttf) — titles / table headers / sub-labels
  KhmerSiemreap   (KhmerOSSiemreap.ttf)   — body text / table cells / footer

Produces:
  generate_report()          — single-page A4 diagnostic PDF (scan + AI result)
  generate_user_report_pdf() — admin user management PDF (bilingual table)
"""

import os
import logging
from io import BytesIO
from datetime import datetime, timezone

logger = logging.getLogger(__name__)

# ── ReportLab ─────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
    Table, TableStyle, HRFlowable,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

PAGE_W, PAGE_H = A4
MARGIN = 1.8 * cm

# ── Colours ───────────────────────────────────────────────────────────────────
BRAND_BLUE  = colors.HexColor("#1A73E8")
BRAND_RED   = colors.HexColor("#DC2626")
BRAND_GREEN = colors.HexColor("#059669")
BRAND_GREY  = colors.HexColor("#64748B")
NAVY        = colors.HexColor("#0F172A")
LIGHT_BG    = colors.HexColor("#F8FAFC")
BORDER_C    = colors.HexColor("#E2E8F0")
WHITE       = colors.white

# ── Cambodian date helpers ────────────────────────────────────────────────────
_KH_MONTHS = ["","មករា","កុម្ភៈ","មីនា","មេសា","ឧសភា","មិថុនា",
               "កក្កដា","សីហា","កញ្ញា","តុលា","វិច្ឆិកា","ធ្នូ"]


def _kh_date_long(dt) -> str:
    """ថ្ងៃ DD ខែ MonthName ឆ្នាំ YYYY"""
    if not dt:
        return "—"
    return f"ថ្ងៃ{dt.day:02d} ខែ{_KH_MONTHS[dt.month]} ឆ្នាំ{dt.year}"


def _kh_date_short(dt) -> str:
    """DD/MM/YYYY"""
    return f"{dt.day:02d}/{dt.month:02d}/{dt.year}" if dt else "—"


def _register_fonts(font_dir: str):
    """Register Battambang TTF (idempotent — safe to call multiple times)."""
    _already = set(pdfmetrics.getRegisteredFontNames())
    reg_path  = os.path.join(font_dir, "Battambang-Regular.ttf")
    bold_path = os.path.join(font_dir, "Battambang-Bold.ttf")

    for name, path in [("Battambang", reg_path), ("Battambang-Bold", bold_path)]:
        if name in _already:
            continue                          # already registered — skip
        if not os.path.isfile(path):
            logger.error("Khmer font not found: %s — Khmer text will show as boxes!", path)
            continue
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            logger.debug("Registered font %s from %s", name, path)
        except Exception as e:
            logger.error("Could not register font %s: %s", name, e)


# ── Khmer OS 3-font system ────────────────────────────────────────────────────
#
#  Slot           Font file                 Glyphs  Latin
#  KhmerMuol    ← KhmerOSMuolLight.ttf     204      ✅  (headers / decorative)
#  KhmerBattambang ← KhmerOSBattambang.ttf  204      ✅  (titles / table headers)
#  KhmerSiemreap ← KhmerOSSiemreap.ttf     139      ❌  (body / cells / footer)
#
#  IMPORTANT: KhmerOSSiemreap contains only Khmer glyphs — no A-Z / a-z.
#  If used directly for mixed Khmer+Latin content (email, names, English
#  status labels) those Latin characters render as □ boxes.
#  register_khmer_fonts() detects this automatically and falls back to
#  Battambang-Regular.ttf (204 glyphs, full Khmer+Latin) for that slot.


def _font_latin_count(ttf_path: str) -> int:
    """Return the number of Latin A-Z / a-z glyphs present in the font."""
    try:
        probe = TTFont("_probe_", ttf_path)
        face  = probe.face
        face.charWidths              # force charToGlyph parse
        return sum(
            1 for c in "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz"
            if face.charToGlyph.get(ord(c)) is not None
        )
    except Exception:
        return 0


def register_khmer_fonts(font_dir: str | None = None) -> None:
    """
    Register the 3 Khmer OS fonts (idempotent).

    KhmerSiemreap falls back to Battambang-Regular.ttf when the Siemreap
    font file is missing Latin characters — otherwise email addresses,
    English names, and status labels like "(Active)" render as □ boxes.
    """
    if font_dir is None:
        font_dir = os.path.normpath(
            os.path.join(os.path.dirname(os.path.abspath(__file__)),
                         "..", "..", "static", "fonts")
        )
    _already = set(pdfmetrics.getRegisteredFontNames())

    # Primary font file for each registered name
    mapping = {
        "KhmerMuol"      : "KhmerOSMuolLight.ttf",
        "KhmerBattambang": "KhmerOSBattambang.ttf",
        "KhmerSiemreap"  : "KhmerOSSiemreap.ttf",
    }

    # Fallback candidates (in priority order) for fonts that lack Latin glyphs.
    # KhmerOSSiemreap.ttf only covers Khmer; use Battambang (full Latin+Khmer)
    # so that email, English names, and status labels render without boxes.
    fallbacks = {
        "KhmerSiemreap": [
            "Battambang-Regular.ttf",   # original font bundled with project
            "KhmerOSBattambang.ttf",    # same font, downloaded copy
        ],
    }

    for font_name, filename in mapping.items():
        if font_name in _already:
            continue

        primary_path = os.path.join(font_dir, filename)
        chosen_path  = None

        if os.path.isfile(primary_path):
            if font_name in fallbacks and _font_latin_count(primary_path) < 26:
                # Primary file lacks Latin glyphs — use the first available fallback
                logger.warning(
                    "Font '%s' (%s) has no Latin glyphs — body text with "
                    "English characters would show as □ boxes. "
                    "Using a fallback font with full Khmer+Latin coverage.",
                    font_name, filename,
                )
                for fb_file in fallbacks[font_name]:
                    fb_path = os.path.join(font_dir, fb_file)
                    if os.path.isfile(fb_path) and _font_latin_count(fb_path) >= 26:
                        chosen_path = fb_path
                        logger.info(
                            "Registered '%s' from fallback: %s", font_name, fb_path
                        )
                        break
                if not chosen_path:
                    # No suitable fallback — use primary and warn
                    chosen_path = primary_path
                    logger.warning(
                        "No Latin-capable fallback found for '%s'. "
                        "English text in body cells may show as □ boxes.",
                        font_name,
                    )
            else:
                chosen_path = primary_path
        else:
            logger.error(
                "Khmer OS font not found: %s — falling back to primary only",
                primary_path,
            )
            # Try fallbacks even when primary is missing
            for fb_file in fallbacks.get(font_name, []):
                fb_path = os.path.join(font_dir, fb_file)
                if os.path.isfile(fb_path):
                    chosen_path = fb_path
                    break
            if not chosen_path:
                continue

        try:
            pdfmetrics.registerFont(TTFont(font_name, chosen_path))
            logger.debug("Registered font '%s' from %s", font_name, chosen_path)
        except Exception as exc:
            logger.error("Could not register font '%s': %s", font_name, exc)


def get_khmer_styles(font_dir: str | None = None) -> dict:
    """
    Register fonts and return a dict of ParagraphStyle objects keyed by role.

    Usage:
        styles = get_khmer_styles()
        Paragraph("Header text", styles['report_header'])
    """
    register_khmer_fonts(font_dir)

    return {
        # ── HEADER styles  (KhmerMuol — decorative thick Khmer) ──────────────
        "report_header": ParagraphStyle(
            "ReportHeader",
            fontName   = "KhmerMuol",
            fontSize   = 20,
            textColor  = colors.HexColor("#6366F1"),
            alignment  = TA_CENTER,
            spaceAfter = 8,
            leading    = 32,
        ),
        "section_header": ParagraphStyle(
            "SectionHeader",
            fontName   = "KhmerMuol",
            fontSize   = 14,
            textColor  = colors.HexColor("#0F172A"),
            alignment  = TA_CENTER,
            spaceAfter = 6,
            leading    = 24,
        ),
        "page_header": ParagraphStyle(
            "PageHeader",
            fontName   = "KhmerMuol",
            fontSize   = 12,
            textColor  = colors.HexColor("#94A3B8"),
            alignment  = TA_RIGHT,
            spaceAfter = 4,
            leading    = 20,
        ),

        # ── TITLE styles  (KhmerBattambang — clean readable Khmer) ───────────
        "report_title": ParagraphStyle(
            "ReportTitle",
            fontName   = "KhmerBattambang",
            fontSize   = 16,
            textColor  = colors.HexColor("#6366F1"),
            alignment  = TA_CENTER,
            spaceAfter = 6,
            leading    = 26,
        ),
        "card_title": ParagraphStyle(
            "CardTitle",
            fontName   = "KhmerBattambang",
            fontSize   = 13,
            textColor  = colors.HexColor("#0F172A"),
            alignment  = TA_LEFT,
            spaceAfter = 4,
            leading    = 22,
        ),
        "table_header": ParagraphStyle(
            "TableHeader",
            fontName   = "KhmerBattambang",
            fontSize   = 11,
            textColor  = colors.white,
            alignment  = TA_CENTER,
            spaceAfter = 0,
            leading    = 18,
        ),
        "sub_title": ParagraphStyle(
            "SubTitle",
            fontName   = "KhmerBattambang",
            fontSize   = 12,
            textColor  = colors.HexColor("#334155"),
            alignment  = TA_CENTER,
            spaceAfter = 4,
            leading    = 20,
        ),

        # ── BODY styles  (KhmerSiemreap — compact readable Khmer) ────────────
        "body_text": ParagraphStyle(
            "BodyText",
            fontName   = "KhmerSiemreap",
            fontSize   = 10,
            textColor  = colors.HexColor("#334155"),
            alignment  = TA_LEFT,
            spaceAfter = 2,
            leading    = 16,
        ),
        "table_cell": ParagraphStyle(
            "TableCell",
            fontName   = "KhmerSiemreap",
            fontSize   = 10,
            textColor  = colors.HexColor("#334155"),
            alignment  = TA_LEFT,
            spaceAfter = 0,
            leading    = 16,
        ),
        "table_cell_center": ParagraphStyle(
            "TableCellCenter",
            fontName   = "KhmerSiemreap",
            fontSize   = 10,
            textColor  = colors.HexColor("#334155"),
            alignment  = TA_CENTER,
            spaceAfter = 0,
            leading    = 16,
        ),
        "footer": ParagraphStyle(
            "Footer",
            fontName   = "KhmerSiemreap",
            fontSize   = 9,
            textColor  = colors.HexColor("#94A3B8"),
            alignment  = TA_CENTER,
            spaceAfter = 0,
            leading    = 14,
        ),
        "status_active": ParagraphStyle(
            "StatusActive",
            fontName   = "KhmerSiemreap",
            fontSize   = 10,
            textColor  = colors.HexColor("#10B981"),
            alignment  = TA_CENTER,
            leading    = 16,
        ),
        "status_suspended": ParagraphStyle(
            "StatusSuspended",
            fontName   = "KhmerSiemreap",
            fontSize   = 10,
            textColor  = colors.HexColor("#EF4444"),
            alignment  = TA_CENTER,
            leading    = 16,
        ),
        "tier_pro": ParagraphStyle(
            "TierPro",
            fontName   = "KhmerBattambang",
            fontSize   = 10,
            textColor  = colors.HexColor("#854D0E"),
            alignment  = TA_CENTER,
            leading    = 16,
        ),
        "tier_free": ParagraphStyle(
            "TierFree",
            fontName   = "KhmerSiemreap",
            fontSize   = 10,
            textColor  = colors.HexColor("#475569"),
            alignment  = TA_CENTER,
            leading    = 16,
        ),
    }


# ── User management PDF ───────────────────────────────────────────────────────

def generate_user_report_pdf(users: list, font_dir: str | None = None) -> BytesIO:
    """
    Build a bilingual (Khmer/English) A4 PDF for the admin user-management report.

    Parameters
    ----------
    users    : list of dicts or ORM User objects with attributes:
               id, username (or full_name), email, tier, is_active,
               scan_count (or scans_today), created_at
    font_dir : optional path override for the fonts directory

    Returns
    -------
    BytesIO  : PDF bytes positioned at offset 0, ready for Response / send_file
    """
    from reportlab.lib.units import mm

    style = get_khmer_styles(font_dir)
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize     = A4,
        rightMargin  = 20 * mm,
        leftMargin   = 20 * mm,
        topMargin    = 20 * mm,
        bottomMargin = 20 * mm,
    )

    elements = []
    now = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")

    # ── Page header (KhmerMuol, right-aligned date) ───────────────────────────
    elements.append(Paragraph(f"កាលបរិច្ឆេទ / Date: {now}", style["page_header"]))
    elements.append(Spacer(1, 4))

    # ── Brand name (English Helvetica) ───────────────────────────────────────
    brand_style = ParagraphStyle(
        "Brand",
        fontName  = "Helvetica-Bold",
        fontSize  = 22,
        textColor = colors.HexColor("#6366F1"),
        alignment = TA_LEFT,
    )
    elements.append(Paragraph("SmartX-Ray", brand_style))
    elements.append(Spacer(1, 4))

    # ── Divider ───────────────────────────────────────────────────────────────
    elements.append(HRFlowable(
        width=     "100%",
        thickness= 2,
        color=     colors.HexColor("#6366F1"),
        spaceAfter=10,
    ))

    # ── Report title (KhmerBattambang) ────────────────────────────────────────
    elements.append(Paragraph("រាយបាការណ៍គ្រប់គ្រងអ្នកប្រើប្រាស់", style["report_title"]))
    elements.append(Paragraph("User Management Report",               style["sub_title"]))

    total = len(users) if hasattr(users, "__len__") else "?"
    elements.append(Paragraph(
        f"<font name='KhmerBattambang' size='10'>"
        f"សរុប / Total: {total} នាក់ / users  |  SmartX-Ray Cambodia AI Platform"
        f"</font>",
        style["body_text"],
    ))
    elements.append(Spacer(1, 12))

    # ── Extra styles for mixed/Latin-only cells ───────────────────────────────
    # Email addresses are always Latin → Helvetica for sharp rendering.
    # Name cells can be Khmer or Latin → KhmerBattambang (204 glyphs, both).
    email_style = ParagraphStyle(
        "EmailCell",
        fontName   = "Helvetica",
        fontSize   = 9,
        textColor  = colors.HexColor("#334155"),
        alignment  = TA_LEFT,
        spaceAfter = 0,
        leading    = 14,
    )
    name_style = ParagraphStyle(
        "NameCell",
        fontName   = "KhmerBattambang",   # 204 glyphs — handles both scripts
        fontSize   = 10,
        textColor  = colors.HexColor("#334155"),
        alignment  = TA_LEFT,
        spaceAfter = 0,
        leading    = 16,
    )
    num_style = ParagraphStyle(
        "NumCell",
        fontName   = "Helvetica",
        fontSize   = 10,
        textColor  = colors.HexColor("#334155"),
        alignment  = TA_CENTER,
        spaceAfter = 0,
        leading    = 16,
    )
    date_style = ParagraphStyle(
        "DateCell",
        fontName   = "Helvetica",
        fontSize   = 9,
        textColor  = colors.HexColor("#64748B"),
        alignment  = TA_CENTER,
        spaceAfter = 0,
        leading    = 14,
    )

    # ── Table headers (KhmerBattambang, white on indigo) ──────────────────────
    headers = [
        Paragraph("លេខ\n#",          style["table_header"]),
        Paragraph("ឈ្មោះ\nName",      style["table_header"]),
        Paragraph("អ៊ីមែល\nEmail",    style["table_header"]),
        Paragraph("កម្រិត\nTier",     style["table_header"]),
        Paragraph("ស្ថានភាព\nStatus", style["table_header"]),
        Paragraph("ស្កែន\nScans",     style["table_header"]),
        Paragraph("ចូលរួម\nJoined",   style["table_header"]),
    ]

    # ── Table rows ────────────────────────────────────────────────────────────
    table_data = [headers]

    def _attr(obj, *keys):
        """Get the first matching attribute from an ORM object or dict."""
        for k in keys:
            if isinstance(obj, dict):
                if k in obj:
                    return obj[k]
            else:
                if hasattr(obj, k):
                    return getattr(obj, k)
        return None

    for u in users:
        uid        = _attr(u, "id")
        username   = _attr(u, "full_name", "username") or "—"
        email      = _attr(u, "email") or "—"
        tier       = _attr(u, "tier") or "free"
        is_active  = _attr(u, "is_active")
        if is_active is None:
            is_active = True
        scan_count = _attr(u, "scan_count", "scans_today") or 0
        created_at = _attr(u, "created_at")
        joined     = str(created_at)[:10] if created_at else "N/A"

        # Tier cell  (KhmerBattambang has full Latin+Khmer)
        if tier == "pro":
            tier_para = Paragraph("Pro *", style["tier_pro"])  # ★ not in font
        else:
            tier_para = Paragraph("ឥតគិតថ្លៃ", style["tier_free"])

        # Status cell — bilingual: Khmer label + English qualifier
        # KhmerSiemreap now falls back to Battambang so Latin renders fine,
        # but using inline rich-text as belt-and-suspenders for the Latin parts.
        if is_active:
            status_para = Paragraph(
                "<font name='KhmerBattambang' size='10'>សកម្ម</font><br/>"
                "<font name='Helvetica' size='9'>(Active)</font>",
                style["status_active"],
            )
        else:
            status_para = Paragraph(
                "<font name='KhmerBattambang' size='10'>ផ្អាក</font><br/>"
                "<font name='Helvetica' size='9'>(Suspended)</font>",
                style["status_suspended"],
            )

        table_data.append([
            Paragraph(str(uid),         num_style),    # always a number
            Paragraph(username,         name_style),   # Khmer or Latin name
            Paragraph(email,            email_style),  # always Latin email
            tier_para,
            status_para,
            Paragraph(str(scan_count),  num_style),    # always a number
            Paragraph(joined,           date_style),   # always YYYY-MM-DD
        ])

    # ── Build table ───────────────────────────────────────────────────────────
    tbl = Table(
        table_data,
        colWidths  = [12*mm, 35*mm, 52*mm, 22*mm, 32*mm, 18*mm, 29*mm],
        repeatRows = 1,
    )
    tbl.setStyle(TableStyle([
        # Header row
        ("BACKGROUND",    (0, 0), (-1,  0), colors.HexColor("#6366F1")),
        ("TEXTCOLOR",     (0, 0), (-1,  0), colors.white),
        ("ALIGN",         (0, 0), (-1,  0), "CENTER"),
        ("FONTNAME",      (0, 0), (-1,  0), "KhmerBattambang"),
        ("FONTSIZE",      (0, 0), (-1,  0), 10),
        ("LINEBELOW",     (0, 0), (-1,  0), 2, colors.HexColor("#4F46E5")),
        # Body rows
        ("FONTNAME",      (0, 1), (-1, -1), "KhmerSiemreap"),
        ("FONTSIZE",      (0, 1), (-1, -1), 10),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F8FAFC")]),
        # Common
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 6),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
    ]))

    elements.append(tbl)
    elements.append(Spacer(1, 16))

    # ── Footer (KhmerSiemreap) ────────────────────────────────────────────────
    elements.append(HRFlowable(
        width=     "100%",
        thickness= 0.5,
        color=     colors.HexColor("#E2E8F0"),
        spaceAfter=6,
    ))
    # Footer: first line is mixed Khmer+Latin, second is pure Latin.
    # Inline font tags ensure correct rendering regardless of KhmerSiemreap
    # fallback status.
    elements.append(Paragraph(
        "<font name='KhmerBattambang' size='9'>"
        "ឯកសារសម្ងាត់ | SmartX-Ray វេទិកា AI វេជ្ជសាស្ត្រ ប្រទេសកម្ពុជា"
        "</font>",
        style["footer"],
    ))
    elements.append(Paragraph(
        "<font name='Helvetica' size='9' color='#94A3B8'>"
        "Confidential — SmartX-Ray Cambodia Medical AI Platform"
        "</font>",
        style["footer"],
    ))

    doc.build(elements)
    buffer.seek(0)
    return buffer


# ── Ad Manager PDF ────────────────────────────────────────────────────────────

def generate_ad_report_pdf(ads: list, lang: str = "en",
                           font_dir: str | None = None) -> BytesIO:
    """
    Build a bilingual (Khmer/English) A4 PDF for the admin Ad Manager report.

    Parameters
    ----------
    ads      : list of Ad ORM objects
    lang     : 'en' | 'km'  — controls label language throughout
    font_dir : optional path override for the fonts directory

    Returns
    -------
    BytesIO positioned at offset 0, ready for Flask Response / send_file
    """
    from reportlab.lib.units import mm

    style  = get_khmer_styles(font_dir)
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize     = A4,
        rightMargin  = 20 * mm,
        leftMargin   = 20 * mm,
        topMargin    = 20 * mm,
        bottomMargin = 20 * mm,
    )

    elements = []
    now = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")

    # ── Date header ───────────────────────────────────────────────────────────
    date_lbl = "កាលបរិច្ឆេទ" if lang == "km" else "Date"
    elements.append(Paragraph(f"{date_lbl}: {now}", style["page_header"]))
    elements.append(Spacer(1, 4))

    # ── Brand name ────────────────────────────────────────────────────────────
    brand_style = ParagraphStyle(
        "AdBrand",
        fontName  = "Helvetica-Bold",
        fontSize  = 22,
        textColor = colors.HexColor("#6366F1"),
        alignment = TA_LEFT,
    )
    elements.append(Paragraph("SmartX-Ray", brand_style))
    elements.append(Spacer(1, 4))
    elements.append(HRFlowable(width="100%", thickness=2,
                               color=colors.HexColor("#6366F1"), spaceAfter=10))

    # ── Report title ──────────────────────────────────────────────────────────
    if lang == "km":
        elements.append(Paragraph("រាយបាការណ៍គ្រប់គ្រងការផ្សាយ",    style["report_title"]))
        elements.append(Paragraph("Ad Manager Report",                  style["sub_title"]))
    else:
        elements.append(Paragraph("Ad Manager Report",                  style["report_title"]))
        elements.append(Paragraph("SmartX-Ray — Cambodia Medical AI Platform", style["sub_title"]))

    # ── Summary stats ─────────────────────────────────────────────────────────
    total        = len(ads)
    active_count = sum(1 for a in ads if getattr(a, "is_active", True))
    total_imp    = sum(getattr(a, "impressions", 0) or 0 for a in ads)
    total_clk    = sum(getattr(a, "clicks",      0) or 0 for a in ads)
    overall_ctr  = round((total_clk / total_imp * 100) if total_imp else 0, 2)

    if lang == "km":
        summary = (f"សរុបការផ្សាយ: {total} | សកម្ម: {active_count} | "
                   f"ការបង្ហាញ: {total_imp:,} | ការចុច: {total_clk:,} | CTR: {overall_ctr}%")
    else:
        summary = (f"Total Ads: {total}  |  Active: {active_count}  |  "
                   f"Impressions: {total_imp:,}  |  Clicks: {total_clk:,}  |  CTR: {overall_ctr}%")

    elements.append(Paragraph(
        f"<font name='KhmerBattambang' size='10'>{summary}</font>",
        style["body_text"],
    ))
    elements.append(Spacer(1, 12))

    # ── Table headers ─────────────────────────────────────────────────────────
    _TH = style["table_header"]
    if lang == "km":
        headers = [
            Paragraph("#",              _TH),
            Paragraph("ចំណងជើង\nTitle", _TH),
            Paragraph("អ្នកផ្សាយ\nAdvertiser", _TH),
            Paragraph("ទីតាំង\nPosition", _TH),
            Paragraph("ការបង្ហាញ\nImpressions", _TH),
            Paragraph("ការចុច\nClicks",   _TH),
            Paragraph("CTR%",             _TH),
            Paragraph("ស្ថានភាព\nStatus", _TH),
            Paragraph("អាទិភាព\nPriority", _TH),
        ]
    else:
        headers = [
            Paragraph("#",             _TH),
            Paragraph("Title",         _TH),
            Paragraph("Advertiser",    _TH),
            Paragraph("Position",      _TH),
            Paragraph("Impressions",   _TH),
            Paragraph("Clicks",        _TH),
            Paragraph("CTR%",          _TH),
            Paragraph("Status",        _TH),
            Paragraph("Priority",      _TH),
        ]

    # ── Extra cell styles ─────────────────────────────────────────────────────
    num_s = ParagraphStyle("AdNum",  fontName="Helvetica",       fontSize=9,
                           textColor=colors.HexColor("#334155"),
                           alignment=TA_CENTER, spaceAfter=0, leading=14)
    txt_s = ParagraphStyle("AdTxt",  fontName="KhmerBattambang", fontSize=9,
                           textColor=colors.HexColor("#334155"),
                           alignment=TA_LEFT,   spaceAfter=0, leading=14)

    PLACEMENT_MAP = {
        "banner":       {"en": "Top Header",    "km": "បដាខាងលើ"},
        "sidebar":      {"en": "Sidebar",       "km": "ចំហៀង"},
        "result_page":  {"en": "Result Page",   "km": "ទំព័រលទ្ធផល"},
        "interstitial": {"en": "Interstitial",  "km": "Interstitial"},
        "footer":       {"en": "Footer",        "km": "ជើងទំព័រ"},
    }

    table_data = [headers]
    for a in ads:
        title_v  = (getattr(a, "title",      None) or "—")[:32]
        adv_v    = (getattr(a, "advertiser", None) or "—")[:22]
        plc_raw  = getattr(a, "placement",   "") or ""
        plc_v    = PLACEMENT_MAP.get(plc_raw, {}).get(lang, plc_raw.replace("_", " ").title())
        imp_v    = getattr(a, "impressions", 0) or 0
        clk_v    = getattr(a, "clicks",      0) or 0
        ctr_v    = round((clk_v / imp_v * 100) if imp_v else 0, 2)
        active_v = getattr(a, "is_active",   True)
        prio_v   = getattr(a, "priority",    0) or 0

        status_txt   = ("សកម្ម" if lang == "km" else "Active") if active_v \
                  else ("អសកម្ម" if lang == "km" else "Inactive")
        status_color = colors.HexColor("#10B981") if active_v else colors.HexColor("#64748B")
        st_s = ParagraphStyle("AdSt", fontName="KhmerBattambang", fontSize=9,
                              textColor=status_color, alignment=TA_CENTER,
                              spaceAfter=0, leading=14)

        table_data.append([
            Paragraph(str(getattr(a, "id", "—")), num_s),
            Paragraph(title_v,  txt_s),
            Paragraph(adv_v,    txt_s),
            Paragraph(plc_v,    txt_s),
            Paragraph(f"{imp_v:,}", num_s),
            Paragraph(f"{clk_v:,}", num_s),
            Paragraph(f"{ctr_v}%",  num_s),
            Paragraph(status_txt,   st_s),
            Paragraph(str(prio_v),  num_s),
        ])

    col_widths = [12, 50, 38, 38, 32, 28, 24, 30, 25]   # mm
    t = Table(table_data, colWidths=[w * mm for w in col_widths], repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1,  0), colors.HexColor("#6366F1")),
        ("TEXTCOLOR",     (0, 0), (-1,  0), colors.white),
        ("FONTNAME",      (0, 0), (-1,  0), "KhmerBattambang"),
        ("FONTSIZE",      (0, 0), (-1,  0), 9),
        ("ALIGN",         (0, 0), (-1,  0), "CENTER"),
        ("LINEBELOW",     (0, 0), (-1,  0), 2, colors.HexColor("#4F46E5")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING",   (0, 0), (-1, -1), 5),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F8FAFC")]),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
    ]))

    elements.append(t)
    elements.append(Spacer(1, 16))

    # ── Footer ────────────────────────────────────────────────────────────────
    elements.append(HRFlowable(width="100%", thickness=0.5,
                               color=colors.HexColor("#E2E8F0"), spaceAfter=6))
    if lang == "km":
        elements.append(Paragraph(
            "<font name='KhmerBattambang' size='9'>"
            "ឯកសារសម្ងាត់ | SmartX-Ray វេទិកា AI វេជ្ជសាស្ត្រ ប្រទេសកម្ពុជា"
            "</font>",
            style["footer"],
        ))
    else:
        elements.append(Paragraph(
            "<font name='Helvetica' size='9'>"
            f"Confidential — SmartX-Ray Cambodia Medical AI Platform • {now}"
            "</font>",
            style["footer"],
        ))

    doc.build(elements)
    buffer.seek(0)
    return buffer


# ── Ad Manager DOCX ───────────────────────────────────────────────────────────

def generate_ad_report_docx(ads: list, lang: str = "en") -> BytesIO:
    """
    Build a bilingual DOCX (Word) report for the admin Ad Manager module.

    Requires the ``python-docx`` package (add to requirements.txt if absent).
    Falls back to a plain-CSV BytesIO if python-docx is not installed so the
    endpoint never returns 500 — just a less-formatted file.

    Parameters
    ----------
    ads  : list of Ad ORM objects
    lang : 'en' | 'km'

    Returns
    -------
    BytesIO positioned at offset 0
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # python-docx not installed — emit plain CSV so download still works
        import csv, io
        out = io.StringIO()
        w = csv.writer(out)
        w.writerow(["#", "Title", "Advertiser", "Placement",
                    "Impressions", "Clicks", "CTR%", "Status", "Priority"])
        for a in ads:
            imp = getattr(a, "impressions", 0) or 0
            clk = getattr(a, "clicks",      0) or 0
            ctr = round((clk / imp * 100) if imp else 0, 2)
            w.writerow([
                getattr(a, "id",         ""),
                getattr(a, "title",      ""),
                getattr(a, "advertiser", ""),
                getattr(a, "placement",  ""),
                imp, clk, f"{ctr}%",
                "Active" if getattr(a, "is_active", True) else "Inactive",
                getattr(a, "priority",   0),
            ])
        buf = BytesIO(out.getvalue().encode("utf-8-sig"))
        buf.seek(0)
        return buf

    now          = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")
    total        = len(ads)
    active_count = sum(1 for a in ads if getattr(a, "is_active", True))
    total_imp    = sum(getattr(a, "impressions", 0) or 0 for a in ads)
    total_clk    = sum(getattr(a, "clicks",      0) or 0 for a in ads)
    overall_ctr  = round((total_clk / total_imp * 100) if total_imp else 0, 2)

    INDIGO = RGBColor(0x63, 0x66, 0xF1)
    GREY   = RGBColor(0x64, 0x74, 0x8B)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Title block ───────────────────────────────────────────────────────────
    brand_p = doc.add_paragraph()
    brand_r = brand_p.add_run("SmartX-Ray")
    brand_r.bold = True
    brand_r.font.size = Pt(20)
    brand_r.font.color.rgb = INDIGO

    if lang == "km":
        main_title = "រាយបាការណ៍គ្រប់គ្រងការផ្សាយ"
        sub_title  = "Ad Manager Report"
    else:
        main_title = "Ad Manager Report"
        sub_title  = "SmartX-Ray — Cambodia Medical AI Platform"

    h = doc.add_heading(main_title, level=1)
    h.runs[0].font.color.rgb = INDIGO

    sub_p  = doc.add_paragraph(sub_title)
    sub_p.runs[0].font.size  = Pt(11)
    sub_p.runs[0].font.color.rgb = GREY

    date_p = doc.add_paragraph()
    date_lbl = "កាលបរិច្ឆេទ" if lang == "km" else "Generated"
    date_p.add_run(f"{date_lbl}: ").bold = True
    date_p.add_run(now)

    doc.add_paragraph()   # spacer

    # ── Summary section ───────────────────────────────────────────────────────
    if lang == "km":
        stats_lines = [
            ("សរុបការផ្សាយ / Total Ads",          str(total)),
            ("សកម្ម / Active",                     str(active_count)),
            ("ការបង្ហាញ / Total Impressions",      f"{total_imp:,}"),
            ("ការចុច / Total Clicks",               f"{total_clk:,}"),
            ("CTR ទូទៅ / Overall CTR",              f"{overall_ctr}%"),
        ]
    else:
        stats_lines = [
            ("Total Ads",         str(total)),
            ("Active",            str(active_count)),
            ("Total Impressions", f"{total_imp:,}"),
            ("Total Clicks",      f"{total_clk:,}"),
            ("Overall CTR",       f"{overall_ctr}%"),
        ]

    for lbl, val in stats_lines:
        sp = doc.add_paragraph()
        sp.add_run(f"{lbl}: ").bold = True
        sp.add_run(val)
        sp.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()   # spacer

    # ── Data table ────────────────────────────────────────────────────────────
    if lang == "km":
        col_labels = ["#", "ចំណងជើង", "អ្នកផ្សាយ", "ទីតាំង",
                      "ការបង្ហាញ", "ការចុច", "CTR%", "ស្ថានភាព", "អាទិភាព"]
    else:
        col_labels = ["#", "Title", "Advertiser", "Position",
                      "Impressions", "Clicks", "CTR%", "Status", "Priority"]

    PLACEMENT_MAP = {
        "banner":       {"en": "Top Header",    "km": "បដាខាងលើ"},
        "sidebar":      {"en": "Sidebar",       "km": "ចំហៀង"},
        "result_page":  {"en": "Result Page",   "km": "ទំព័រលទ្ធផល"},
        "interstitial": {"en": "Interstitial",  "km": "Interstitial"},
        "footer":       {"en": "Footer",        "km": "ជើងទំព័រ"},
    }

    table = doc.add_table(rows=1, cols=len(col_labels))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, lbl in enumerate(col_labels):
        hdr_cells[i].text = lbl
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Indigo background via cell XML shading
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "6366F1")
        tc_pr.append(shd)

    # Data rows
    for a in ads:
        plc_raw  = getattr(a, "placement", "") or ""
        plc_v    = PLACEMENT_MAP.get(plc_raw, {}).get(lang, plc_raw.replace("_", " ").title())
        imp_v    = getattr(a, "impressions", 0) or 0
        clk_v    = getattr(a, "clicks",      0) or 0
        ctr_v    = round((clk_v / imp_v * 100) if imp_v else 0, 2)
        active_v = getattr(a, "is_active",   True)

        status_v = ("សកម្ម" if lang == "km" else "Active") if active_v \
              else ("អសកម្ម" if lang == "km" else "Inactive")

        row_cells = table.add_row().cells
        values = [
            str(getattr(a, "id",         "—")),
            getattr(a, "title",           "") or "—",
            getattr(a, "advertiser",      "") or "—",
            plc_v,
            f"{imp_v:,}",
            f"{clk_v:,}",
            f"{ctr_v}%",
            status_v,
            str(getattr(a, "priority",   0)),
        ]
        for i, val in enumerate(values):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_r = foot_p.add_run(
        f"{'ឯកសារសម្ងាត់ | ' if lang == 'km' else 'Confidential — '}"
        f"SmartX-Ray Cambodia Medical AI Platform • {now}"
    )
    foot_r.font.size      = Pt(8)
    foot_r.font.color.rgb = GREY
    foot_p.alignment      = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Subscription Report PDF ───────────────────────────────────────────────────

def generate_subscription_report_pdf(subs: list, lang: str = "en",
                                     font_dir: str | None = None) -> BytesIO:
    """
    Build a bilingual (Khmer/English) A4 PDF for the admin Subscriptions report.

    Parameters
    ----------
    subs     : list of dicts with keys:
               id, user_name, user_email, plan, status,
               period_end (datetime|None), cancel_at_end (bool), created_at (datetime|None)
    lang     : 'en' | 'km'
    font_dir : optional font-directory override

    Returns
    -------
    BytesIO positioned at offset 0
    """
    from reportlab.lib.units import mm

    style  = get_khmer_styles(font_dir)
    buffer = BytesIO()
    now    = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )
    elements = []

    # ── Date header ────────────────────────────────────────────────────────────
    date_lbl = "កាលបរិច្ឆេទ" if lang == "km" else "Date"
    elements.append(Paragraph(f"{date_lbl}: {now}", style["page_header"]))
    elements.append(Spacer(1, 4))

    # ── Brand ──────────────────────────────────────────────────────────────────
    brand_sty = ParagraphStyle("SubBrand", fontName="Helvetica-Bold", fontSize=22,
                               textColor=colors.HexColor("#6366F1"), alignment=TA_LEFT)
    elements.append(Paragraph("SmartX-Ray", brand_sty))
    elements.append(Spacer(1, 4))
    elements.append(HRFlowable(width="100%", thickness=2,
                               color=colors.HexColor("#6366F1"), spaceAfter=10))

    # ── Title ──────────────────────────────────────────────────────────────────
    if lang == "km":
        elements.append(Paragraph("រាយបាការណ៍ការជាវ",          style["report_title"]))
        elements.append(Paragraph("Subscription Report",          style["sub_title"]))
    else:
        elements.append(Paragraph("Subscription Report",          style["report_title"]))
        elements.append(Paragraph("SmartX-Ray — ABA KHQR Billing", style["sub_title"]))

    # ── Revenue summary ─────────────────────────────────────────────────────────
    total         = len(subs)
    active_monthly  = sum(1 for s in subs if s.get("status") == "active"   and s.get("plan") == "monthly")
    active_yearly   = sum(1 for s in subs if s.get("status") == "active"   and s.get("plan") == "yearly")
    trialing        = sum(1 for s in subs if s.get("status") == "trialing")
    mrr             = round(active_monthly * 9.99,  2)
    arr             = round(active_yearly  * 79.99, 2)

    if lang == "km":
        summary = (f"ការជាវសរុប: {total}  |  ប្រចាំខែ: {active_monthly} (${mrr:.2f} MRR)"
                   f"  |  ប្រចាំឆ្នាំ: {active_yearly} (${arr:.2f} ARR)"
                   f"  |  សាកល្បង: {trialing}")
    else:
        summary = (f"Total: {total}  |  Monthly: {active_monthly} (${mrr:.2f} MRR)"
                   f"  |  Yearly: {active_yearly} (${arr:.2f} ARR)"
                   f"  |  Trialing: {trialing}")

    elements.append(Paragraph(
        f"<font name='KhmerBattambang' size='10'>{summary}</font>",
        style["body_text"],
    ))
    elements.append(Spacer(1, 12))

    # ── Table headers ─────────────────────────────────────────────────────────
    _TH = style["table_header"]
    if lang == "km":
        headers = [
            Paragraph("#",                _TH),
            Paragraph("ឈ្មោះ\nName",      _TH),
            Paragraph("អ៊ីមែល\nEmail",    _TH),
            Paragraph("គម្រោង\nPlan",     _TH),
            Paragraph("ស្ថានភាព\nStatus", _TH),
            Paragraph("ផ្ការ៉ូស\nRenewal", _TH),
            Paragraph("បន្តស្វ័យ\nAuto-Renew", _TH),
            Paragraph("ចាប់ពី\nSince",    _TH),
        ]
    else:
        headers = [
            Paragraph("#",           _TH),
            Paragraph("Name",        _TH),
            Paragraph("Email",       _TH),
            Paragraph("Plan",        _TH),
            Paragraph("Status",      _TH),
            Paragraph("Renewal",     _TH),
            Paragraph("Auto-Renew",  _TH),
            Paragraph("Since",       _TH),
        ]

    # ── Cell styles ────────────────────────────────────────────────────────────
    num_s  = ParagraphStyle("SubNum",   fontName="Helvetica",       fontSize=9,
                            textColor=colors.HexColor("#334155"),
                            alignment=TA_CENTER, spaceAfter=0, leading=14)
    txt_s  = ParagraphStyle("SubTxt",   fontName="KhmerBattambang", fontSize=9,
                            textColor=colors.HexColor("#334155"),
                            alignment=TA_LEFT, spaceAfter=0, leading=14)
    email_s = ParagraphStyle("SubMail", fontName="Helvetica",       fontSize=8,
                             textColor=colors.HexColor("#475569"),
                             alignment=TA_LEFT, spaceAfter=0, leading=13)

    STATUS_MAP = {
        "active":   {"en": "Active",    "km": "សកម្ម",      "color": "#059669"},
        "trialing": {"en": "Trial",     "km": "សាកល្បង",    "color": "#D97706"},
        "past_due": {"en": "Past Due",  "km": "ហួសកំណត់",   "color": "#DC2626"},
        "canceled": {"en": "Cancelled", "km": "បានបោះបង់",  "color": "#64748B"},
        "inactive": {"en": "Inactive",  "km": "អសកម្ម",     "color": "#64748B"},
    }

    table_data = [headers]
    for s in subs:
        uid     = s.get("id", "—")
        name    = (s.get("user_name")  or "—")[:28]
        email   = (s.get("user_email") or "—")
        plan    = s.get("plan", "")
        status  = s.get("status", "")
        pe      = s.get("period_end")
        ca      = s.get("created_at")
        cancel  = s.get("cancel_at_end", False)

        plan_lbl   = ("ប្រចាំខែ" if lang == "km" else "Monthly") if plan == "monthly" \
                else ("ប្រចាំឆ្នាំ ★" if lang == "km" else "Yearly ★")
        status_info = STATUS_MAP.get(status, {"en": status, "km": status, "color": "#64748B"})
        status_lbl  = status_info.get(lang, status_info["en"])
        st_s = ParagraphStyle("SubSt", fontName="KhmerBattambang", fontSize=9,
                              textColor=colors.HexColor(status_info["color"]),
                              alignment=TA_CENTER, spaceAfter=0, leading=14)

        renewal_str = str(pe)[:10] if pe else "—"
        since_str   = str(ca)[:10] if ca else "—"
        renew_lbl   = ("ទេ" if lang == "km" else "Off") if cancel \
                 else ("បាទ" if lang == "km" else "On")

        table_data.append([
            Paragraph(str(uid),       num_s),
            Paragraph(name,           txt_s),
            Paragraph(email,          email_s),
            Paragraph(plan_lbl,       txt_s),
            Paragraph(status_lbl,     st_s),
            Paragraph(renewal_str,    num_s),
            Paragraph(renew_lbl,      num_s),
            Paragraph(since_str,      num_s),
        ])

    col_widths = [12, 40, 52, 28, 28, 30, 28, 28]   # mm
    t = Table(table_data, colWidths=[w * mm for w in col_widths], repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1,  0), colors.HexColor("#6366F1")),
        ("TEXTCOLOR",     (0, 0), (-1,  0), colors.white),
        ("FONTNAME",      (0, 0), (-1,  0), "KhmerBattambang"),
        ("FONTSIZE",      (0, 0), (-1,  0), 9),
        ("ALIGN",         (0, 0), (-1,  0), "CENTER"),
        ("LINEBELOW",     (0, 0), (-1,  0), 2, colors.HexColor("#4F46E5")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING",   (0, 0), (-1, -1), 5),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F8FAFC")]),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 16))

    # ── Footer ─────────────────────────────────────────────────────────────────
    elements.append(HRFlowable(width="100%", thickness=0.5,
                               color=colors.HexColor("#E2E8F0"), spaceAfter=6))
    footer_txt = (
        f"<font name='KhmerBattambang' size='9'>"
        f"{'ឯកសារសម្ងាត់ | SmartX-Ray ABA KHQR Billing' if lang == 'km' else 'Confidential — SmartX-Ray ABA KHQR Billing'}"
        f" • {now}</font>"
    )
    elements.append(Paragraph(footer_txt, style["footer"]))

    doc.build(elements)
    buffer.seek(0)
    return buffer


# ── Subscription Report DOCX ──────────────────────────────────────────────────

def generate_subscription_report_docx(subs: list, lang: str = "en") -> BytesIO:
    """
    Build a bilingual DOCX (Word) report for the admin Subscriptions module.

    Falls back to UTF-8 CSV if python-docx is not installed.

    Parameters
    ----------
    subs : list of dicts (same shape as generate_subscription_report_pdf)
    lang : 'en' | 'km'

    Returns
    -------
    BytesIO positioned at offset 0
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        import csv, io
        out = io.StringIO()
        w   = csv.writer(out)
        w.writerow(["#","Name","Email","Plan","Status","Renewal","Auto-Renew","Since"])
        for s in subs:
            w.writerow([
                s.get("id",""),  s.get("user_name",""),  s.get("user_email",""),
                s.get("plan",""), s.get("status",""),
                str(s.get("period_end",""))[:10],
                "Off" if s.get("cancel_at_end") else "On",
                str(s.get("created_at",""))[:10],
            ])
        buf = BytesIO(out.getvalue().encode("utf-8-sig"))
        buf.seek(0)
        return buf

    now    = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")
    INDIGO = RGBColor(0x63, 0x66, 0xF1)
    GREY   = RGBColor(0x64, 0x74, 0x8B)

    # ── Revenue stats ──────────────────────────────────────────────────────────
    total         = len(subs)
    act_monthly   = sum(1 for s in subs if s.get("status") == "active" and s.get("plan") == "monthly")
    act_yearly    = sum(1 for s in subs if s.get("status") == "active" and s.get("plan") == "yearly")
    trialing      = sum(1 for s in subs if s.get("status") == "trialing")
    mrr           = round(act_monthly * 9.99,  2)
    arr           = round(act_yearly  * 79.99, 2)
    total_rev     = round(mrr + arr, 2)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Brand ──────────────────────────────────────────────────────────────────
    brand_p = doc.add_paragraph()
    brand_r = brand_p.add_run("SmartX-Ray")
    brand_r.bold = True;  brand_r.font.size = Pt(20);  brand_r.font.color.rgb = INDIGO

    if lang == "km":
        h = doc.add_heading("រាយបាការណ៍ការជាវ", level=1)
    else:
        h = doc.add_heading("Subscription Report", level=1)
    h.runs[0].font.color.rgb = INDIGO

    sub_p = doc.add_paragraph("SmartX-Ray — ABA KHQR Billing")
    sub_p.runs[0].font.size = Pt(11); sub_p.runs[0].font.color.rgb = GREY

    date_p = doc.add_paragraph()
    date_p.add_run(f"{'កាលបរិច្ឆេទ' if lang == 'km' else 'Generated'}: ").bold = True
    date_p.add_run(now)
    doc.add_paragraph()

    # ── Revenue summary ─────────────────────────────────────────────────────────
    if lang == "km":
        stats_lines = [
            ("ការជាវសរុប / Total",           str(total)),
            ("ប្រចាំខែ / Monthly",            f"{act_monthly}  (${mrr:.2f} MRR)"),
            ("ប្រចាំឆ្នាំ / Yearly",           f"{act_yearly}  (${arr:.2f} ARR)"),
            ("ចំណូលសរុប / Total Revenue",     f"${total_rev:.2f}"),
            ("សាកល្បង / Trialing",            str(trialing)),
        ]
    else:
        stats_lines = [
            ("Total Subscriptions", str(total)),
            ("Monthly Active",      f"{act_monthly}  (${mrr:.2f} MRR)"),
            ("Yearly Active",       f"{act_yearly}  (${arr:.2f} ARR)"),
            ("Total Revenue",       f"${total_rev:.2f}"),
            ("Trialing",            str(trialing)),
        ]

    for lbl, val in stats_lines:
        sp = doc.add_paragraph();  sp.paragraph_format.space_after = Pt(2)
        sp.add_run(f"{lbl}: ").bold = True
        sp.add_run(val)

    doc.add_paragraph()

    # ── Table ──────────────────────────────────────────────────────────────────
    if lang == "km":
        col_labels = ["#","ឈ្មោះ","អ៊ីមែល","គម្រោង","ស្ថានភាព","ផ្ការ៉ូស","បន្តស្វ័យ","ចាប់ពី"]
    else:
        col_labels = ["#","Name","Email","Plan","Status","Renewal","Auto-Renew","Since"]

    STATUS_MAP = {
        "active":   {"en":"Active",    "km":"សកម្ម"},
        "trialing": {"en":"Trial",     "km":"សាកល្បង"},
        "past_due": {"en":"Past Due",  "km":"ហួសកំណត់"},
        "canceled": {"en":"Cancelled", "km":"បានបោះបង់"},
        "inactive": {"en":"Inactive",  "km":"អសកម្ម"},
    }

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _shade_cell(cell, hex_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    table = doc.add_table(rows=1, cols=len(col_labels))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, lbl in enumerate(col_labels):
        hdr_cells[i].text = lbl
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True;  run.font.size = Pt(9);  run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _shade_cell(hdr_cells[i], "6366F1")

    for s in subs:
        plan   = s.get("plan", "")
        status = s.get("status", "")
        pe     = s.get("period_end")
        ca     = s.get("created_at")
        cancel = s.get("cancel_at_end", False)

        plan_lbl   = ("ប្រចាំខែ" if lang == "km" else "Monthly") if plan == "monthly" \
                else ("ប្រចាំឆ្នាំ ★" if lang == "km" else "Yearly ★")
        info       = STATUS_MAP.get(status, {"en": status, "km": status})
        status_lbl = info.get(lang, info["en"])
        renew_lbl  = ("ទេ" if lang == "km" else "Off") if cancel \
                else ("បាទ" if lang == "km" else "On")

        row_cells = table.add_row().cells
        values = [
            str(s.get("id","—")),
            s.get("user_name","") or "—",
            s.get("user_email","") or "—",
            plan_lbl, status_lbl,
            str(pe)[:10] if pe else "—",
            renew_lbl,
            str(ca)[:10] if ca else "—",
        ]
        for i, val in enumerate(values):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # ── Footer ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_r = foot_p.add_run(
        f"{'ឯកសារសម្ងាត់ | ' if lang == 'km' else 'Confidential — '}"
        f"SmartX-Ray ABA KHQR Billing • {now}"
    )
    foot_r.font.size = Pt(8);  foot_r.font.color.rgb = GREY
    foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Doctor Marketplace PDF ────────────────────────────────────────────────────

def generate_doctor_report_pdf(doctors: list, lang: str = "en",
                               font_dir: str | None = None) -> BytesIO:
    """
    Build a bilingual (Khmer/English) A4 PDF for the admin Doctor Marketplace report.

    Parameters
    ----------
    doctors  : list of Doctor ORM objects
    lang     : 'en' | 'km'
    font_dir : optional font-directory override

    Returns
    -------
    BytesIO positioned at offset 0
    """
    from reportlab.lib.units import mm

    style  = get_khmer_styles(font_dir)
    buffer = BytesIO()
    now    = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")

    doc = SimpleDocTemplate(
        buffer,
        pagesize     = A4,
        rightMargin  = 18 * mm,
        leftMargin   = 18 * mm,
        topMargin    = 20 * mm,
        bottomMargin = 20 * mm,
    )
    elements = []

    # ── Date header ───────────────────────────────────────────────────────────
    date_lbl = "កាលបរិច្ឆេទ" if lang == "km" else "Date"
    elements.append(Paragraph(f"{date_lbl}: {now}", style["page_header"]))
    elements.append(Spacer(1, 4))

    # ── Brand ─────────────────────────────────────────────────────────────────
    brand_sty = ParagraphStyle("DocBrand", fontName="Helvetica-Bold", fontSize=22,
                               textColor=colors.HexColor("#6366F1"), alignment=TA_LEFT)
    elements.append(Paragraph("SmartX-Ray", brand_sty))
    elements.append(Spacer(1, 4))
    elements.append(HRFlowable(width="100%", thickness=2,
                               color=colors.HexColor("#6366F1"), spaceAfter=10))

    # ── Title ─────────────────────────────────────────────────────────────────
    if lang == "km":
        elements.append(Paragraph("រាយបាការណ៍ផ្សារវេជ្ជបណ្ឌិត",    style["report_title"]))
        elements.append(Paragraph("Doctor Marketplace Report",          style["sub_title"]))
    else:
        elements.append(Paragraph("Doctor Marketplace Report",          style["report_title"]))
        elements.append(Paragraph("SmartX-Ray — Cambodia Medical AI Platform", style["sub_title"]))

    # ── Summary stats ─────────────────────────────────────────────────────────
    total    = len(doctors)
    approved = sum(1 for d in doctors if getattr(d, "is_verified", False) and getattr(d, "is_active", True))
    pending  = sum(1 for d in doctors if not getattr(d, "is_verified", False) and getattr(d, "is_active", True))
    rejected = sum(1 for d in doctors if not getattr(d, "is_active", True))
    featured = sum(1 for d in doctors if getattr(d, "is_featured", False))

    if lang == "km":
        summary = (f"វេជ្ជបណ្ឌិតសរុប: {total}  |  អនុម័ត: {approved}  |  "
                   f"រង់ចាំ: {pending}  |  បដិសេធ: {rejected}  |  ពិសេស: {featured}")
    else:
        summary = (f"Total: {total}  |  Approved: {approved}  |  "
                   f"Pending: {pending}  |  Rejected: {rejected}  |  Featured: {featured}")

    elements.append(Paragraph(
        f"<font name='KhmerBattambang' size='10'>{summary}</font>",
        style["body_text"],
    ))
    elements.append(Spacer(1, 12))

    # ── Table headers ─────────────────────────────────────────────────────────
    _TH = style["table_header"]
    if lang == "km":
        headers = [
            Paragraph("#",                    _TH),
            Paragraph("ឈ្មោះ\nName",          _TH),
            Paragraph("ឯកទេស\nSpecialty",     _TH),
            Paragraph("មន្ទីរពេទ្យ\nHospital", _TH),
            Paragraph("ទីក្រុង\nCity",         _TH),
            Paragraph("ការវាយ\nRating",        _TH),
            Paragraph("ស្ថានភាព\nStatus",      _TH),
            Paragraph("ពិសេស\nFeatured",       _TH),
            Paragraph("ចាប់ពី\nSince",         _TH),
        ]
    else:
        headers = [
            Paragraph("#",          _TH),
            Paragraph("Name",       _TH),
            Paragraph("Specialty",  _TH),
            Paragraph("Hospital",   _TH),
            Paragraph("City",       _TH),
            Paragraph("Rating",     _TH),
            Paragraph("Status",     _TH),
            Paragraph("Featured",   _TH),
            Paragraph("Since",      _TH),
        ]

    # ── Cell styles ───────────────────────────────────────────────────────────
    num_s   = ParagraphStyle("DocNum",  fontName="Helvetica",       fontSize=9,
                             textColor=colors.HexColor("#334155"),
                             alignment=TA_CENTER, spaceAfter=0, leading=14)
    txt_s   = ParagraphStyle("DocTxt",  fontName="KhmerBattambang", fontSize=9,
                             textColor=colors.HexColor("#334155"),
                             alignment=TA_LEFT,   spaceAfter=0, leading=14)
    email_s = ParagraphStyle("DocMail", fontName="Helvetica",       fontSize=8,
                             textColor=colors.HexColor("#475569"),
                             alignment=TA_LEFT,   spaceAfter=0, leading=13)

    STATUS_MAP = {
        "approved": {"en": "Verified",  "km": "បានផ្ទៀងផ្ទាត់", "color": "#059669"},
        "pending":  {"en": "Pending",   "km": "រង់ចាំ",           "color": "#D97706"},
        "rejected": {"en": "Rejected",  "km": "បានបដិសេធ",        "color": "#DC2626"},
    }

    def _doc_status_str(d):
        if getattr(d, "is_verified", False) and getattr(d, "is_active", True):
            return "approved"
        if not getattr(d, "is_active", True):
            return "rejected"
        return "pending"

    table_data = [headers]
    for doc in doctors:
        name_v  = (getattr(doc, "full_name",     None) or "—")[:30]
        sp_v    = (getattr(doc, "specialty",     None) or "—")[:22]
        hosp_v  = (getattr(doc, "hospital",      None) or "—")[:22]
        city_v  = (getattr(doc, "city",          None) or "—")[:16]
        rating  = getattr(doc, "rating",         0) or 0
        rev_cnt = getattr(doc, "review_count",   0) or 0
        feat_v  = getattr(doc, "is_featured",    False)
        ca      = getattr(doc, "created_at",     None)

        st      = _doc_status_str(doc)
        st_info = STATUS_MAP[st]
        st_lbl  = st_info.get(lang, st_info["en"])
        st_sty  = ParagraphStyle(f"DocSt{st}", fontName="KhmerBattambang", fontSize=9,
                                 textColor=colors.HexColor(st_info["color"]),
                                 alignment=TA_CENTER, spaceAfter=0, leading=14)

        feat_lbl = ("ពិសេស" if lang == "km" else "★ Yes") if feat_v \
              else ("ធម្មតា"  if lang == "km" else "No")
        feat_sty = ParagraphStyle("DocFeat", fontName="KhmerBattambang", fontSize=9,
                                  textColor=colors.HexColor("#78350F" if feat_v else "#64748B"),
                                  alignment=TA_CENTER, spaceAfter=0, leading=14)

        rating_str = f"{float(rating):.1f} ({rev_cnt})" if rating else "—"
        since_str  = str(ca)[:10] if ca else "—"

        table_data.append([
            Paragraph(str(getattr(doc, "id", "—")), num_s),
            Paragraph(name_v,      txt_s),
            Paragraph(sp_v,        txt_s),
            Paragraph(hosp_v,      txt_s),
            Paragraph(city_v,      txt_s),
            Paragraph(rating_str,  num_s),
            Paragraph(st_lbl,      st_sty),
            Paragraph(feat_lbl,    feat_sty),
            Paragraph(since_str,   num_s),
        ])

    col_widths = [12, 40, 36, 36, 28, 26, 30, 24, 24]   # mm  (total ≈ 256 mm ≤ A4 usable)
    t = Table(table_data, colWidths=[w * mm for w in col_widths], repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1,  0), colors.HexColor("#6366F1")),
        ("TEXTCOLOR",     (0, 0), (-1,  0), colors.white),
        ("FONTNAME",      (0, 0), (-1,  0), "KhmerBattambang"),
        ("FONTSIZE",      (0, 0), (-1,  0), 9),
        ("ALIGN",         (0, 0), (-1,  0), "CENTER"),
        ("LINEBELOW",     (0, 0), (-1,  0), 2, colors.HexColor("#4F46E5")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING",   (0, 0), (-1, -1), 5),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1),
         [colors.white, colors.HexColor("#F8FAFC")]),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8F0")),
    ]))

    elements.append(t)
    elements.append(Spacer(1, 16))

    # ── Footer ────────────────────────────────────────────────────────────────
    elements.append(HRFlowable(width="100%", thickness=0.5,
                               color=colors.HexColor("#E2E8F0"), spaceAfter=6))
    if lang == "km":
        elements.append(Paragraph(
            "<font name='KhmerBattambang' size='9'>"
            "ឯកសារសម្ងាត់ | SmartX-Ray ផ្សារវេជ្ជបណ្ឌិតកម្ពុជា"
            "</font>",
            style["footer"],
        ))
    else:
        elements.append(Paragraph(
            f"<font name='Helvetica' size='9'>"
            f"Confidential — SmartX-Ray Cambodia Doctor Marketplace • {now}"
            f"</font>",
            style["footer"],
        ))

    doc.build(elements)
    buffer.seek(0)
    return buffer


# ── Doctor Marketplace DOCX ───────────────────────────────────────────────────

def generate_doctor_report_docx(doctors: list, lang: str = "en") -> BytesIO:
    """
    Build a bilingual DOCX (Word) report for the admin Doctor Marketplace module.

    Falls back to UTF-8 CSV if python-docx is not installed.

    Parameters
    ----------
    doctors : list of Doctor ORM objects
    lang    : 'en' | 'km'

    Returns
    -------
    BytesIO positioned at offset 0
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        # python-docx not installed — emit plain CSV
        import csv, io
        out = io.StringIO()
        w   = csv.writer(out)
        w.writerow(["#", "Name", "Specialty", "Hospital", "City",
                    "Rating", "Status", "Featured", "Since"])
        for d in doctors:
            is_ver  = getattr(d, "is_verified", False)
            is_act  = getattr(d, "is_active",   True)
            st      = "Approved" if (is_ver and is_act) else ("Rejected" if not is_act else "Pending")
            rating  = getattr(d, "rating", 0) or 0
            w.writerow([
                getattr(d, "id",          ""),
                getattr(d, "full_name",   ""),
                getattr(d, "specialty",   ""),
                getattr(d, "hospital",    ""),
                getattr(d, "city",        ""),
                f"{float(rating):.1f}" if rating else "—",
                st,
                "Yes" if getattr(d, "is_featured", False) else "No",
                str(getattr(d, "created_at", ""))[:10],
            ])
        buf = BytesIO(out.getvalue().encode("utf-8-sig"))
        buf.seek(0)
        return buf

    now    = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")
    total  = len(doctors)
    approved = sum(1 for d in doctors if getattr(d, "is_verified", False) and getattr(d, "is_active", True))
    pending  = sum(1 for d in doctors if not getattr(d, "is_verified", False) and getattr(d, "is_active", True))
    rejected = sum(1 for d in doctors if not getattr(d, "is_active", True))
    featured = sum(1 for d in doctors if getattr(d, "is_featured", False))

    INDIGO = RGBColor(0x63, 0x66, 0xF1)
    GREY   = RGBColor(0x64, 0x74, 0x8B)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Brand ─────────────────────────────────────────────────────────────────
    brand_p = doc.add_paragraph()
    brand_r = brand_p.add_run("SmartX-Ray")
    brand_r.bold = True; brand_r.font.size = Pt(20); brand_r.font.color.rgb = INDIGO

    if lang == "km":
        h = doc.add_heading("រាយបាការណ៍ផ្សារវេជ្ជបណ្ឌិត", level=1)
    else:
        h = doc.add_heading("Doctor Marketplace Report", level=1)
    h.runs[0].font.color.rgb = INDIGO

    sub_p = doc.add_paragraph(
        "SmartX-Ray — Cambodia Doctor Marketplace"
    )
    sub_p.runs[0].font.size = Pt(11); sub_p.runs[0].font.color.rgb = GREY

    date_p = doc.add_paragraph()
    date_p.add_run(f"{'កាលបរិច្ឆេទ' if lang == 'km' else 'Generated'}: ").bold = True
    date_p.add_run(now)
    doc.add_paragraph()

    # ── Summary ───────────────────────────────────────────────────────────────
    if lang == "km":
        stats_lines = [
            ("វេជ្ជបណ្ឌិតសរុប / Total",   str(total)),
            ("អនុម័ត / Approved",           str(approved)),
            ("រង់ចាំ / Pending",            str(pending)),
            ("បដិសេធ / Rejected",           str(rejected)),
            ("ពិសេស / Featured",            str(featured)),
        ]
    else:
        stats_lines = [
            ("Total Doctors",  str(total)),
            ("Approved",       str(approved)),
            ("Pending",        str(pending)),
            ("Rejected",       str(rejected)),
            ("Featured",       str(featured)),
        ]

    for lbl, val in stats_lines:
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)
        sp.add_run(f"{lbl}: ").bold = True
        sp.add_run(val)

    doc.add_paragraph()

    # ── Table ─────────────────────────────────────────────────────────────────
    if lang == "km":
        col_labels = ["#", "ឈ្មោះ", "ឯកទេស", "មន្ទីរពេទ្យ",
                      "ទីក្រុង", "ការវាយ", "ស្ថានភាព", "ពិសេស", "ចាប់ពី"]
    else:
        col_labels = ["#", "Name", "Specialty", "Hospital",
                      "City", "Rating", "Status", "Featured", "Since"]

    STATUS_MAP = {
        "approved": {"en": "Verified",  "km": "បានផ្ទៀងផ្ទាត់"},
        "pending":  {"en": "Pending",   "km": "រង់ចាំ"},
        "rejected": {"en": "Rejected",  "km": "បានបដិសេធ"},
    }

    def _shade_cell(cell, hex_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tc_pr.append(shd)

    def _doc_st(d):
        if getattr(d, "is_verified", False) and getattr(d, "is_active", True):
            return "approved"
        if not getattr(d, "is_active", True):
            return "rejected"
        return "pending"

    table = doc.add_table(rows=1, cols=len(col_labels))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, lbl in enumerate(col_labels):
        hdr_cells[i].text = lbl
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr_cells[i], "6366F1")

    for d in doctors:
        st       = _doc_st(d)
        info     = STATUS_MAP[st]
        st_lbl   = info.get(lang, info["en"])
        rating   = getattr(d, "rating", 0) or 0
        rev_cnt  = getattr(d, "review_count", 0) or 0
        feat_v   = getattr(d, "is_featured", False)
        ca       = getattr(d, "created_at", None)

        feat_lbl = ("ពិសេស" if lang == "km" else "★ Yes") if feat_v \
              else ("ធម្មតា"  if lang == "km" else "No")

        row_cells = table.add_row().cells
        values = [
            str(getattr(d, "id",          "—")),
            getattr(d, "full_name",        "—") or "—",
            getattr(d, "specialty",        "—") or "—",
            getattr(d, "hospital",         "—") or "—",
            getattr(d, "city",             "—") or "—",
            f"{float(rating):.1f} ({rev_cnt})" if rating else "—",
            st_lbl,
            feat_lbl,
            str(ca)[:10] if ca else "—",
        ]
        for i, val in enumerate(values):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_r = foot_p.add_run(
        f"{'ឯកសារសម្ងាត់ | ' if lang == 'km' else 'Confidential — '}"
        f"SmartX-Ray Cambodia Doctor Marketplace • {now}"
    )
    foot_r.font.size      = Pt(8)
    foot_r.font.color.rgb = GREY
    foot_p.alignment      = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── System Logs PDF ───────────────────────────────────────────────────────────

def generate_log_report_pdf(logs: list, lang: str = "en",
                            font_dir: str | None = None) -> BytesIO:
    """
    Build a bilingual (Khmer/English) A4 PDF for the admin System Logs report.

    Parameters
    ----------
    logs     : list of SystemLog ORM objects
    lang     : 'en' | 'km'
    font_dir : optional font-directory override

    Returns
    -------
    BytesIO positioned at offset 0
    """
    from reportlab.lib.units import mm

    style  = get_khmer_styles(font_dir)
    buffer = BytesIO()
    now    = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")

    doc = SimpleDocTemplate(
        buffer,
        pagesize     = A4,
        rightMargin  = 16 * mm,
        leftMargin   = 16 * mm,
        topMargin    = 20 * mm,
        bottomMargin = 20 * mm,
    )
    elements = []

    # ── Date header ───────────────────────────────────────────────────────────
    date_lbl = "កាលបរិច្ឆេទ" if lang == "km" else "Date"
    elements.append(Paragraph(f"{date_lbl}: {now}", style["page_header"]))
    elements.append(Spacer(1, 4))

    # ── Brand ─────────────────────────────────────────────────────────────────
    brand_sty = ParagraphStyle("LogBrand", fontName="Helvetica-Bold", fontSize=22,
                               textColor=colors.HexColor("#6366F1"), alignment=TA_LEFT)
    elements.append(Paragraph("SmartX-Ray", brand_sty))
    elements.append(Spacer(1, 4))
    elements.append(HRFlowable(width="100%", thickness=2,
                               color=colors.HexColor("#6366F1"), spaceAfter=10))

    # ── Title ─────────────────────────────────────────────────────────────────
    if lang == "km":
        elements.append(Paragraph("រាយបាការណ៍កំណត់ហេតុប្រព័ន្ធ", style["report_title"]))
        elements.append(Paragraph("System Logs Report",               style["sub_title"]))
    else:
        elements.append(Paragraph("System Logs Report",               style["report_title"]))
        elements.append(Paragraph("SmartX-Ray — Cambodia Medical AI Platform", style["sub_title"]))

    # ── Summary ───────────────────────────────────────────────────────────────
    total    = len(logs)
    critical = sum(1 for l in logs if getattr(l, "severity", "") == "critical")
    high     = sum(1 for l in logs if getattr(l, "severity", "") == "high")
    warning  = sum(1 for l in logs if getattr(l, "severity", "") == "warning")
    info_cnt = sum(1 for l in logs if getattr(l, "severity", "") == "info")

    if lang == "km":
        summary = (f"សរុប: {total}  |  វិបត្តិ: {critical}  |  "
                   f"ខ្ពស់: {high}  |  ការព្រមាន: {warning}  |  ព័ត៌មាន: {info_cnt}")
    else:
        summary = (f"Total: {total}  |  Critical: {critical}  |  "
                   f"High: {high}  |  Warning: {warning}  |  Info: {info_cnt}")

    elements.append(Paragraph(
        f"<font name='KhmerBattambang' size='10'>{summary}</font>",
        style["body_text"],
    ))
    elements.append(Spacer(1, 10))

    # ── Table headers ─────────────────────────────────────────────────────────
    _TH = style["table_header"]
    if lang == "km":
        headers = [
            Paragraph("#",                    _TH),
            Paragraph("ពេល\nTime",            _TH),
            Paragraph("ព្រឹត្តិការណ៍\nEvent", _TH),
            Paragraph("ភាព\nSeverity",        _TH),
            Paragraph("អ្នកប្រើ\nUser",        _TH),
            Paragraph("សារ\nMessage",         _TH),
            Paragraph("AI(ms)",               _TH),
            Paragraph("IP",                   _TH),
        ]
    else:
        headers = [
            Paragraph("#",          _TH),
            Paragraph("Time",       _TH),
            Paragraph("Event",      _TH),
            Paragraph("Severity",   _TH),
            Paragraph("User",       _TH),
            Paragraph("Message",    _TH),
            Paragraph("AI (ms)",    _TH),
            Paragraph("IP",         _TH),
        ]

    # ── Cell styles ───────────────────────────────────────────────────────────
    num_s   = ParagraphStyle("LNum",  fontName="Helvetica",       fontSize=8,
                             textColor=colors.HexColor("#64748B"),
                             alignment=TA_CENTER, spaceAfter=0, leading=12)
    ts_s    = ParagraphStyle("LTs",   fontName="Helvetica",       fontSize=7.5,
                             textColor=colors.HexColor("#64748B"),
                             alignment=TA_LEFT,   spaceAfter=0, leading=12)
    txt_s   = ParagraphStyle("LTxt",  fontName="KhmerBattambang", fontSize=8.5,
                             textColor=colors.HexColor("#334155"),
                             alignment=TA_LEFT,   spaceAfter=0, leading=13)
    msg_s   = ParagraphStyle("LMsg",  fontName="KhmerBattambang", fontSize=8,
                             textColor=colors.HexColor("#334155"),
                             alignment=TA_LEFT,   spaceAfter=0, leading=12)
    ip_s    = ParagraphStyle("LIp",   fontName="Helvetica",       fontSize=7.5,
                             textColor=colors.HexColor("#64748B"),
                             alignment=TA_CENTER, spaceAfter=0, leading=12)

    SEV_COLORS = {
        "info":     "#1D4ED8",
        "warning":  "#92400E",
        "high":     "#9A3412",
        "critical": "#991B1B",
    }
    SEV_BG = {
        "info":     "#EFF6FF",
        "warning":  "#FEF3C7",
        "high":     "#FFEDD5",
        "critical": "#FEE2E2",
    }
    SEV_KM = {
        "info": "ព័ត៌មាន", "warning": "ការព្រមាន",
        "high": "ខ្ពស់",    "critical": "វិបត្តិ",
    }
    EVT_KM = {
        "scan": "ស្កែន", "auth_login": "ចូលប្រព័ន្ធ",
        "auth_fail": "Login បរាជ័យ", "admin_action": "Admin",
        "telegram_alert": "Telegram", "error": "កំហុស",
        "health_check": "សុខភាព",
    }

    table_data = [headers]
    for l in logs:
        lid  = getattr(l, "id",            "—")
        ca   = getattr(l, "created_at",    None)
        ts   = str(ca)[:19].replace("T", " ") if ca else "—"
        evt  = getattr(l, "event_type",    "") or ""
        sev  = getattr(l, "severity",      "info") or "info"
        email= getattr(l, "user_email",    None)
        if email is None and hasattr(l, "user") and l.user:
            email = l.user.email
        msg  = (getattr(l, "message",      "") or "")[:80]
        ms   = getattr(l, "processing_ms", None)
        ip   = getattr(l, "ip_address",    None) or "—"

        sev_lbl = (SEV_KM.get(sev, sev) if lang == "km" else sev.capitalize())
        evt_lbl = (EVT_KM.get(evt, evt) if lang == "km" else evt.replace("_", " ").title())
        ms_str  = f"{ms:,}ms" if ms else "—"

        sev_col = SEV_COLORS.get(sev, "#64748B")
        sev_bg  = SEV_BG.get(sev, "#F8FAFC")
        sev_sty = ParagraphStyle(f"Sev{sev}", fontName="KhmerBattambang", fontSize=8,
                                 textColor=colors.HexColor(sev_col),
                                 alignment=TA_CENTER, spaceAfter=0, leading=12)

        table_data.append([
            Paragraph(str(lid),  num_s),
            Paragraph(ts,        ts_s),
            Paragraph(evt_lbl,   txt_s),
            Paragraph(sev_lbl,   sev_sty),
            Paragraph(email or "—", msg_s),
            Paragraph(msg,       msg_s),
            Paragraph(ms_str,    num_s),
            Paragraph(ip,        ip_s),
        ])

    col_widths = [14, 38, 36, 28, 42, 68, 24, 26]   # mm  (≈ 276 mm total on A4)
    t = Table(table_data, colWidths=[w * mm for w in col_widths], repeatRows=1)

    # Row background from severity
    sev_cmd = []
    for i, l in enumerate(logs, start=1):
        bg = SEV_BG.get(getattr(l, "severity", "info"), "#FFFFFF")
        sev_cmd.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor(bg)))

    t.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1,  0), colors.HexColor("#6366F1")),
        ("TEXTCOLOR",     (0, 0), (-1,  0), colors.white),
        ("FONTNAME",      (0, 0), (-1,  0), "KhmerBattambang"),
        ("FONTSIZE",      (0, 0), (-1,  0), 8),
        ("ALIGN",         (0, 0), (-1,  0), "CENTER"),
        ("LINEBELOW",     (0, 0), (-1,  0), 1.5, colors.HexColor("#4F46E5")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 4),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
        ("GRID",          (0, 0), (-1, -1), 0.4, colors.HexColor("#E2E8F0")),
    ] + sev_cmd))

    elements.append(t)
    elements.append(Spacer(1, 14))

    # ── Footer ────────────────────────────────────────────────────────────────
    elements.append(HRFlowable(width="100%", thickness=0.5,
                               color=colors.HexColor("#E2E8F0"), spaceAfter=6))
    if lang == "km":
        elements.append(Paragraph(
            "<font name='KhmerBattambang' size='9'>"
            "ឯកសារសម្ងាត់ | SmartX-Ray ប្រព័ន្ធ Log"
            "</font>",
            style["footer"],
        ))
    else:
        elements.append(Paragraph(
            f"<font name='Helvetica' size='9'>"
            f"Confidential — SmartX-Ray System Logs • {now}"
            f"</font>",
            style["footer"],
        ))

    doc.build(elements)
    buffer.seek(0)
    return buffer


# ── System Logs DOCX ──────────────────────────────────────────────────────────

def generate_log_report_docx(logs: list, lang: str = "en") -> BytesIO:
    """
    Build a bilingual DOCX (Word) report for the admin System Logs module.

    Falls back to UTF-8 CSV if python-docx is not installed.

    Parameters
    ----------
    logs : list of SystemLog ORM objects
    lang : 'en' | 'km'

    Returns
    -------
    BytesIO positioned at offset 0
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        import csv, io
        out = io.StringIO()
        w   = csv.writer(out)
        w.writerow(["#", "Time", "Event", "Severity", "User", "Message", "AI(ms)", "IP"])
        for l in logs:
            ca = getattr(l, "created_at", None)
            w.writerow([
                getattr(l, "id", ""),
                str(ca)[:19].replace("T", " ") if ca else "",
                getattr(l, "event_type",    ""),
                getattr(l, "severity",      ""),
                getattr(l, "user_email",    "") or "",
                getattr(l, "message",       "") or "",
                getattr(l, "processing_ms", "") or "",
                getattr(l, "ip_address",    "") or "",
            ])
        buf = BytesIO(out.getvalue().encode("utf-8-sig"))
        buf.seek(0)
        return buf

    now    = datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")
    total  = len(logs)
    critical = sum(1 for l in logs if getattr(l, "severity", "") == "critical")
    high     = sum(1 for l in logs if getattr(l, "severity", "") == "high")
    warning  = sum(1 for l in logs if getattr(l, "severity", "") == "warning")
    info_cnt = sum(1 for l in logs if getattr(l, "severity", "") == "info")

    INDIGO = RGBColor(0x63, 0x66, 0xF1)
    GREY   = RGBColor(0x64, 0x74, 0x8B)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Brand ─────────────────────────────────────────────────────────────────
    brand_p = doc.add_paragraph()
    brand_r = brand_p.add_run("SmartX-Ray")
    brand_r.bold = True; brand_r.font.size = Pt(20); brand_r.font.color.rgb = INDIGO

    if lang == "km":
        h = doc.add_heading("រាយបាការណ៍កំណត់ហេតុប្រព័ន្ធ", level=1)
    else:
        h = doc.add_heading("System Logs Report", level=1)
    h.runs[0].font.color.rgb = INDIGO

    sub_p = doc.add_paragraph("SmartX-Ray — Cambodia Medical AI Platform")
    sub_p.runs[0].font.size = Pt(11); sub_p.runs[0].font.color.rgb = GREY

    date_p = doc.add_paragraph()
    date_p.add_run(f"{'កាលបរិច្ឆេទ' if lang == 'km' else 'Generated'}: ").bold = True
    date_p.add_run(now)
    doc.add_paragraph()

    # ── Summary ───────────────────────────────────────────────────────────────
    if lang == "km":
        stats_lines = [
            ("សរុប / Total",          str(total)),
            ("វិបត្តិ / Critical",    str(critical)),
            ("ខ្ពស់ / High",           str(high)),
            ("ការព្រមាន / Warning",   str(warning)),
            ("ព័ត៌មាន / Info",         str(info_cnt)),
        ]
    else:
        stats_lines = [
            ("Total Logs",  str(total)),
            ("Critical",    str(critical)),
            ("High",        str(high)),
            ("Warning",     str(warning)),
            ("Info",        str(info_cnt)),
        ]

    for lbl, val in stats_lines:
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)
        sp.add_run(f"{lbl}: ").bold = True
        sp.add_run(val)

    doc.add_paragraph()

    # ── Table ─────────────────────────────────────────────────────────────────
    if lang == "km":
        col_labels = ["#", "ពេល", "ព្រឹត្តិការណ៍", "ភាព", "អ្នកប្រើ", "សារ", "AI(ms)", "IP"]
    else:
        col_labels = ["#", "Time", "Event", "Severity", "User", "Message", "AI(ms)", "IP"]

    SEV_FILL = {
        "info":     "EFF6FF", "warning": "FEF3C7",
        "high":     "FFEDD5", "critical": "FEE2E2",
    }
    SEV_KM = {
        "info": "ព័ត៌មាន", "warning": "ការព្រមាន",
        "high": "ខ្ពស់",    "critical": "វិបត្តិ",
    }
    EVT_KM = {
        "scan": "ស្កែន", "auth_login": "ចូលប្រព័ន្ធ",
        "auth_fail": "Login បរាជ័យ", "admin_action": "Admin",
        "telegram_alert": "Telegram", "error": "កំហុស",
        "health_check": "សុខភាព",
    }

    def _shade(cell, hex_color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tc_pr.append(shd)

    table = doc.add_table(rows=1, cols=len(col_labels))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, lbl in enumerate(col_labels):
        hdr_cells[i].text = lbl
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True; run.font.size = Pt(8); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _shade(hdr_cells[i], "6366F1")

    for l in logs:
        ca    = getattr(l, "created_at", None)
        sev   = getattr(l, "severity",   "info") or "info"
        evt   = getattr(l, "event_type", "") or ""
        email = getattr(l, "user_email", None)
        if email is None and hasattr(l, "user") and l.user:
            email = l.user.email
        ms    = getattr(l, "processing_ms", None)

        sev_lbl = (SEV_KM.get(sev, sev) if lang == "km" else sev.capitalize())
        evt_lbl = (EVT_KM.get(evt, evt) if lang == "km" else evt.replace("_", " ").title())

        row_cells = table.add_row().cells
        values = [
            str(getattr(l, "id",       "—")),
            str(ca)[:19].replace("T", " ") if ca else "—",
            evt_lbl,
            sev_lbl,
            email or "—",
            (getattr(l, "message", "") or "")[:80],
            f"{ms:,}ms" if ms else "—",
            getattr(l, "ip_address", "") or "—",
        ]
        fill = SEV_FILL.get(sev, "FFFFFF")
        for i, val in enumerate(values):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(8)
            if sev in ("critical", "high", "warning"):
                _shade(row_cells[i], fill)

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    foot_p = doc.add_paragraph()
    foot_r = foot_p.add_run(
        f"{'ឯកសារសម្ងាត់ | ' if lang == 'km' else 'Confidential — '}"
        f"SmartX-Ray System Logs • {now}"
    )
    foot_r.font.size      = Pt(8)
    foot_r.font.color.rgb = GREY
    foot_p.alignment      = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Main entry ────────────────────────────────────────────────────────────────

def generate_report(scan, user, output_dir: str) -> tuple[str, int]:
    """
    Build the bilingual (Khmer/English) PDF report for a single scan.

    Parameters
    ----------
    scan       : Scan ORM object
    user       : User ORM object
    output_dir : absolute path to static/reports/

    Returns
    -------
    (filename, file_size_bytes)
    """
    from flask import current_app

    font_dir = os.path.normpath(
        os.path.join(current_app.root_path, "..", "static", "fonts")
    )
    _register_fonts(font_dir)

    S = _StyleFactory()

    filename = (
        f"report_scan{scan.id}_"
        f"{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}.pdf"
    )
    filepath = os.path.join(output_dir, filename)

    doc = SimpleDocTemplate(
        filepath, pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=MARGIN,  bottomMargin=MARGIN,
    )

    story = []

    # ── 1. Letterhead ─────────────────────────────────────────────────────────
    now_utc  = datetime.now(timezone.utc)
    date_kh  = _kh_date_long(scan.created_at)

    lh = Table([[
        Paragraph("SmartX-Ray", S("brand")),
        Paragraph(
            f"<font name='Battambang' size='8'>"
            f"<b>ទីក្រុងភ្នំពេញ, ព្រះរាជាណាចក្រកម្ពុជា</b></font><br/>"
            f"<font name='Helvetica' size='7.5' color='#64748B'>"
            f"Phnom Penh, Kingdom of Cambodia</font>",
            S("lh_right"),
        ),
    ]], colWidths=[100*cm/10, PAGE_W - 2*MARGIN - 100*cm/10])
    lh.setStyle(TableStyle([
        ("VALIGN",  (0,0),(-1,-1), "MIDDLE"),
        ("PADDING", (0,0),(-1,-1), 0),
    ]))
    story += [
        lh,
        Spacer(1, 0.25*cm),
        HRFlowable(width="100%", thickness=2, color=BRAND_BLUE, spaceAfter=0.3*cm),
    ]

    # Bilingual title
    story += [
        Paragraph("របាយការណ៍ត្រួតពិនិត្យរូបថតអ៊ិចរ៉េ AI", S("title_kh")),
        Paragraph("AI-Powered Pneumonia Detection Report",   S("title_en")),
        Spacer(1, 0.4*cm),
    ]

    # ── 2. Patient & scan info ────────────────────────────────────────────────
    def bi(km, en):
        """Bilingual label cell."""
        return Paragraph(
            f"<font name='Battambang-Bold' size='8' color='#64748B'>{km}</font><br/>"
            f"<font name='Helvetica'       size='7' color='#94a3b8'>{en}</font>",
            S("info_label"),
        )

    def val(text):
        return Paragraph(
            f"<font name='Battambang' size='9'>{text}</font>",
            S("info_val"),
        )

    info = Table([
        [bi("ឈ្មោះអ្នកជំងឺ","Patient Name"),
         val(user.full_name or "—"),
         bi("កាលបរិច្ឆេទស្កែន","Scan Date"),
         val(date_kh)],
        [bi("អ៊ីមែល","Email"),
         val(user.email or "—"),
         bi("លេខស្កែន","Scan ID"),
         val(f"#{scan.id}")],
        [bi("កំណែម៉ូដែល","Model Version"),
         val(scan.model_version or "v1.0"),
         bi("បង្កើតរបាយការណ៍","Report Generated"),
         val(_kh_date_short(now_utc))],
    ], colWidths=[3.8*cm, 7.2*cm, 3.8*cm, 3.8*cm])

    info.setStyle(TableStyle([
        ("ROWBACKGROUNDS", (0,0),(-1,-1), [WHITE, LIGHT_BG]),
        ("GRID",     (0,0),(-1,-1), 0.3, BORDER_C),
        ("PADDING",  (0,0),(-1,-1), 6),
        ("VALIGN",   (0,0),(-1,-1), "MIDDLE"),
    ]))
    story += [info, Spacer(1, 0.6*cm)]

    # ── 3. Diagnosis result ───────────────────────────────────────────────────
    is_pneumonia   = scan.prediction == "PNEUMONIA"
    result_color   = BRAND_RED if is_pneumonia else BRAND_GREEN
    result_kh      = "ជំងឺរលាកសួត"   if is_pneumonia else "ធម្មតា"
    result_en      = "PNEUMONIA"       if is_pneumonia else "NORMAL"
    confidence_pct = round(scan.confidence * 100, 2)

    story += [
        Paragraph(
            "<font name='Battambang-Bold' size='11' color='#1A73E8'>លទ្ធផលវិនិច្ឆ័យ</font>"
            "  <font name='Helvetica' size='9' color='#64748B'>/ Diagnosis Result</font>",
            S("section_head"),
        ),
        Spacer(1, 0.2*cm),
    ]

    # Result box
    result_box = Table([[
        Paragraph(
            f"<font name='Battambang-Bold' size='26' color='{result_color.hexval()}'>"
            f"{result_kh}</font><br/>"
            f"<font name='Helvetica-Bold' size='14' color='{result_color.hexval()}'>"
            f"{result_en}</font>",
            S("result_main"),
        ),
        Paragraph(
            f"<font name='Battambang' size='9' color='#64748B'>ភាពជឿជាក់ / Confidence</font><br/>"
            f"<font name='Helvetica-Bold' size='28' color='{result_color.hexval()}'>"
            f"{confidence_pct}%</font>",
            S("result_conf"),
        ),
    ]], colWidths=[(PAGE_W-2*MARGIN)*0.55, (PAGE_W-2*MARGIN)*0.45])
    result_box.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(-1,-1), colors.HexColor(
            "#FEF2F2" if is_pneumonia else "#ECFDF5")),
        ("PADDING",    (0,0),(-1,-1), 14),
        ("VALIGN",     (0,0),(-1,-1), "MIDDLE"),
        ("LINEBELOW",  (0,0),(-1,0),  2, result_color),
        ("BOX",        (0,0),(-1,-1), 1, BORDER_C),
    ]))
    story += [result_box, Spacer(1, 0.3*cm)]

    # Confidence bar
    bar_w     = PAGE_W - 2*MARGIN
    filled_w  = bar_w * (confidence_pct / 100)
    empty_w   = bar_w - filled_w
    conf_bar  = Table([["", ""]], colWidths=[filled_w, empty_w], rowHeights=[10])
    conf_bar.setStyle(TableStyle([
        ("BACKGROUND", (0,0),(0,0), result_color),
        ("BACKGROUND", (1,0),(1,0), BORDER_C),
        ("PADDING",    (0,0),(-1,-1), 0),
    ]))
    story += [conf_bar, Spacer(1, 0.6*cm)]

    # ── 4. Clinical note ──────────────────────────────────────────────────────
    if is_pneumonia:
        note_kh = (
            "ម៉ូដែល AI បានរកឃើញលំនាំដែលស្របនឹង <b>ជំងឺរលាកសួត</b> "
            "នៅក្នុងរូបថតអ៊ិចរ៉េទ្រូងដែលបានផ្ទុកឡើង។ "
            "នេះជាការរកឃើញដោយ AI ហើយត្រូវតែត្រូវបានពិនិត្យឡើងវិញ "
            "ដោយវិទ្យុវិទ្យាករ ឬគ្រូពេទ្យដែលមានលក្ខណៈសម្បត្តិ "
            "មុននឹងធ្វើការសម្រេចចិត្តព្យាបាល។"
        )
        note_en = (
            "The AI model detected patterns consistent with <b>pneumonia</b> "
            "in the uploaded chest X-ray. This is an AI-assisted finding and "
            "must be reviewed by a qualified radiologist or physician before "
            "any clinical decision is made."
        )
    else:
        note_kh = (
            "ម៉ូដែល AI <b>មិនបានរកឃើញលំនាំដូចជំងឺរលាកសួតគួរឱ្យកត់សម្គាល់</b> "
            "នៅក្នុងរូបថតអ៊ិចរ៉េទ្រូង។ "
            "នេះមិនមែនជាការបញ្ជាក់ពីសុខភាពផ្លូវវេជ្ជសាស្ត្រទេ។ "
            "សូមពិគ្រោះជាមួយគ្រូពេទ្យប្រសិនបើអ្នកមានរោគសញ្ញា។"
        )
        note_en = (
            "The AI model found <b>no significant pneumonia-like patterns</b> "
            "in the uploaded chest X-ray. This does not constitute a medical "
            "clearance. Please consult a physician if you have clinical symptoms."
        )

    story += [
        Paragraph(
            "<font name='Battambang-Bold' size='11' color='#1A73E8'>កំណត់ចំណាំព្យាបាល</font>"
            "  <font name='Helvetica' size='9' color='#64748B'>/ Clinical Note</font>",
            S("section_head"),
        ),
        Spacer(1, 0.15*cm),
        Paragraph(
            f"<font name='Battambang' size='9'>{note_kh}</font>",
            S("body_kh"),
        ),
        Spacer(1, 0.1*cm),
        Paragraph(
            f"<font name='Helvetica' size='8.5' color='#475569'>{note_en}</font>",
            S("body_en"),
        ),
        Spacer(1, 0.6*cm),
    ]

    # ── 5. Images ─────────────────────────────────────────────────────────────
    static_dir = os.path.join(current_app.root_path, "..", "static")
    orig_path  = (os.path.join(static_dir, scan.image_path)
                  if scan.image_path else None)
    heat_path  = (os.path.join(static_dir, scan.heatmap_path)
                  if scan.heatmap_path else None)
    img_w = (PAGE_W - 2*MARGIN - 1*cm) / 2

    def img_cell(path, cap_kh, cap_en):
        if path and os.path.exists(path):
            img = RLImage(path, width=img_w, height=img_w * 0.85)
            cap = Paragraph(
                f"<font name='Battambang' size='8' color='#64748B'>{cap_kh}</font><br/>"
                f"<font name='Helvetica'   size='7' color='#94a3b8'>{cap_en}</font>",
                S("caption"),
            )
            return [img, cap]
        return [
            Paragraph(
                f"<font name='Battambang' size='8' color='#94a3b8'>"
                f"<i>{cap_kh} — មិនអាចរកបាន</i></font>",
                S("caption"),
            ),
            "",
        ]

    xray_cell  = img_cell(orig_path,  "រូបថតអ៊ិចរ៉េដើម",   "Original X-Ray")
    heatmap_cell = img_cell(heat_path, "ផែនទីកំដៅ Grad-CAM", "Grad-CAM Heatmap")

    story += [
        Paragraph(
            "<font name='Battambang-Bold' size='11' color='#1A73E8'>រូបភាព</font>"
            "  <font name='Helvetica' size='9' color='#64748B'>/ Imaging</font>",
            S("section_head"),
        ),
        Spacer(1, 0.2*cm),
        Table(
            [[xray_cell[0], heatmap_cell[0]],
             [xray_cell[1], heatmap_cell[1]]],
            colWidths=[img_w, img_w],
            hAlign="LEFT",
        ),
        Spacer(1, 0.7*cm),
    ]

    # ── 6. Disclaimer ─────────────────────────────────────────────────────────
    story += [
        HRFlowable(width="100%", thickness=0.5, color=BORDER_C),
        Spacer(1, 0.2*cm),
        Paragraph(
            "<font name='Battambang-Bold' size='8'>គំហើញ:</font>"
            "<font name='Battambang' size='8'> SmartX-Ray គឺជាឧបករណ៍ជំនួយ AI "
            "សម្រាប់វិជ្ជាជីវៈពេទ្យប៉ុណ្ណោះ។ "
            "វាមិនអាចជំនួសដំបូន្មានវេជ្ជសាស្ត្រ ការធ្វើរោគវិនិច្ឆ័យ "
            "ឬការព្យាបាលពីអ្នកជំនាញបានទេ។ "
            "សូមតែងតែពិគ្រោះជាមួយអ្នកផ្តល់សុខភាពដែលមានគុណភាព។</font>",
            S("disclaimer"),
        ),
        Spacer(1, 0.1*cm),
        Paragraph(
            "<font name='Helvetica-Bold' size='7.5'>Disclaimer:</font>"
            "<font name='Helvetica' size='7.5'> SmartX-Ray is a computer-aided "
            "detection tool for trained medical professionals only. It is NOT a "
            "substitute for professional medical advice, diagnosis, or treatment. "
            "Always seek guidance from a qualified health provider.</font>",
            S("disclaimer"),
        ),
    ]

    doc.build(story)
    file_size = os.path.getsize(filepath)
    logger.info("PDF report generated: %s (%d bytes)", filepath, file_size)
    return filename, file_size


# ── Style factory ─────────────────────────────────────────────────────────────

class _StyleFactory:
    """Returns ParagraphStyle objects by name (creates lazily)."""

    def __init__(self):
        self._base = getSampleStyleSheet()
        self._cache: dict[str, ParagraphStyle] = {}

    def __call__(self, name: str) -> ParagraphStyle:
        if name not in self._cache:
            self._cache[name] = self._make(name)
        return self._cache[name]

    def _make(self, name: str) -> ParagraphStyle:
        b = self._base["Normal"]
        defs = {
            "brand": dict(fontName="Helvetica-Bold", fontSize=22,
                          textColor=BRAND_BLUE, spaceAfter=0),
            "lh_right": dict(fontName="Battambang", fontSize=8,
                             textColor=BRAND_GREY, alignment=TA_RIGHT, leading=13),
            "title_kh": dict(fontName="Battambang-Bold", fontSize=17,
                             textColor=NAVY, alignment=TA_CENTER,
                             spaceAfter=2, leading=24),
            "title_en": dict(fontName="Helvetica", fontSize=10,
                             textColor=BRAND_GREY, alignment=TA_CENTER, spaceAfter=0),
            "section_head": dict(fontName="Battambang-Bold", fontSize=11,
                                 textColor=BRAND_BLUE, spaceBefore=4, spaceAfter=3),
            "info_label": dict(fontName="Battambang-Bold", fontSize=8,
                               textColor=BRAND_GREY, leading=12),
            "info_val":   dict(fontName="Battambang", fontSize=9,
                               textColor=NAVY, leading=13),
            "result_main": dict(fontName="Battambang-Bold", fontSize=26,
                                alignment=TA_LEFT, leading=32),
            "result_conf": dict(fontName="Helvetica-Bold", fontSize=28,
                                alignment=TA_CENTER, leading=34),
            "body_kh":    dict(fontName="Battambang", fontSize=9,
                               textColor=NAVY, leading=15),
            "body_en":    dict(fontName="Helvetica",  fontSize=8.5,
                               textColor=colors.HexColor("#475569"), leading=13),
            "caption":    dict(fontName="Battambang", fontSize=8,
                               textColor=BRAND_GREY, alignment=TA_CENTER, leading=12),
            "disclaimer": dict(fontName="Battambang", fontSize=8,
                               textColor=BRAND_GREY, leading=12),
        }
        kwargs = defs.get(name, {})
        return ParagraphStyle(name, parent=b, **kwargs)
